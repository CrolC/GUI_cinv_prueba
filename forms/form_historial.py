import customtkinter as ctk
from tkinter import messagebox
from fpdf import FPDF
from docx import Document
import pandas as pd
import sys
sys.path.append('d:/Python_Proyectos/INTER_C3')


class FormHistorial(ctk.CTk):
    
    def __init__(self):
        super().__init__()

        self.title("Generar Reporte")
        self.geometry("300x200")

        # Botón con ícono de libro
        self.generate_report_button = ctk.CTkButton(self, text="Generar Reporte", 
                                                    command=self.generate_report,
                                                    image=ctk.CTkImage("D:\Python_Proyectos\INTER_C3\imagenes\icon_reporte.png"))
        self.generate_report_button.pack(pady=20)

    def generate_report(self):
        # Aquí puedes cambiar el formato a PDF, DOC o Excel según necesites
        file_format = "PDF"  # Cambia a "DOC", "Excel" según el formato que quieras

        if file_format == "PDF":
            self.generate_pdf_report()
        elif file_format == "DOC":
            self.generate_doc_report()
        elif file_format == "Excel":
            self.generate_excel_report()

    def generate_pdf_report(self):
        #PDF (ejemplo)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Reporte Generado", ln=True, align="C")
        pdf.output("reporte.pdf")
        messagebox.showinfo("Reporte", "PDF generado exitosamente.")

    def generate_doc_report(self):
        # DOC (info ejemplo)
        doc = Document()
        doc.add_heading('Reporte Generado', 0)
        doc.add_paragraph('Este es un reporte generado en formato Word.')
        doc.save("reporte.docx")
        messagebox.showinfo("Reporte", "Documento DOC generado exitosamente.")

    def generate_excel_report(self):
        # Excel (ejem)
        data = {'Columna 1': [1, 2, 3], 'Columna 2': [4, 5, 6]}
        df = pd.DataFrame(data)
        df.to_excel("reporte.xlsx", index=False)
        messagebox.showinfo("Reporte", "Archivo Excel generado exitosamente.")

